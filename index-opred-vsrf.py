import os
import ssl
import hashlib
import asyncio
import re
from pathlib import Path
from bs4 import BeautifulSoup
from aiohttp import ClientSession, TCPConnector
from docx import Document

# === НАСТРОЙКИ ===
BASE_DIR = r"Y:\documents_base\zakonodatelstvo\sud_practica_vushih_sudov"
HASH_FILE = "downloaded_hashes.txt"
ERROR_404_FILE = "404_urls.txt"
BASE_URL = "https://legalacts.ru/sud/"
RETRY_COUNT = 3
RETRY_DELAY = 15        # 15 сек между повторами при ошибках
PAGE_DELAY = 60        # 2 минуты между страницами
DOC_DELAY = 5           # 5 сек между документами

ssl_context = ssl._create_unverified_context()
HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/122.0.0.0 Safari/537.36"
}


def load_hashes(path):
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return set(line.strip() for line in f)
    return set()


def save_hashes(path, hashes):
    with open(path, "w", encoding="utf-8") as f:
        for h in hashes:
            f.write(f"{h}\n")


def save_404(url):
    with open(ERROR_404_FILE, "a", encoding="utf-8") as f:
        f.write(url + "\n")


def hash_content(text: str) -> str:
    return hashlib.md5(text.encode("utf-8")).hexdigest()


async def fetch(session, url, retries=RETRY_COUNT):
    for attempt in range(retries):
        try:
            async with session.get(url, ssl=ssl_context, timeout=30) as response:
                if response.status == 200:
                    return await response.text()
                elif response.status == 404:
                    print(f"❌ HTTP 404: {url}")
                    save_404(url)
                    return None
                elif response.status in (403, 503):
                    print(f"⚠️ HTTP {response.status} (попытка {attempt + 1}/{retries}): {url}")
                else:
                    print(f"⚠️ HTTP {response.status} при загрузке {url}")
        except Exception as e:
            print(f"❌ Ошибка при загрузке {url}: {e}")
        await asyncio.sleep(RETRY_DELAY)
    return None


async def parse_document(session, url, existing_hashes):
    html = await fetch(session, url)
    if not html:
        return

    soup = BeautifulSoup(html, "html.parser")
    meta_title = soup.find("meta", {"property": "og:title"})
    if not meta_title:
        print(f"❌ Не найден заголовок: {url}")
        return

    title = meta_title.get("content", "Документ").replace('\n', ' ').strip()
    safe_title = re.sub(r'[\\/*?:"<>|]', '', title)[:140]
    filename = safe_title + ".docx"
    filepath = os.path.join(BASE_DIR, filename)

    # Попробуем найти текст документа по универсальному классу
    content_block = soup.find("div", class_="main-center-block")
    if not content_block:
        print(f"⚠️ Не найден текст документа: {url}")
        return

    text_parts = [p.get_text(strip=True) for p in content_block.find_all("p")]
    full_text = "\n".join([t for t in text_parts if t])
    if not full_text.strip():
        print(f"⚠️ Пустой документ: {url}")
        return

    content_hash = hash_content(full_text)
    if content_hash in existing_hashes:
        print(f"⏭ Уже скачан по содержимому: {url}")
        return

    doc = Document()
    doc.add_heading(title, 0)
    for paragraph in text_parts:
        if paragraph:
            doc.add_paragraph(paragraph)

    Path(BASE_DIR).mkdir(parents=True, exist_ok=True)
    doc.save(filepath)
    print(f"✅ Сохранён: {filename}")

    existing_hashes.add(content_hash)
    save_hashes(HASH_FILE, existing_hashes)
    await asyncio.sleep(DOC_DELAY)


async def parse_page(session, page_number, existing_hashes):
    url = f"{BASE_URL}?page={page_number}"
    html = await fetch(session, url)
    if not html:
        return False

    soup = BeautifulSoup(html, "html.parser")
    links = soup.select("a[href^='/sud/opredelenie']")
    if not links:
        return False

    for link in links:
        href = link.get("href")
        if href:
            full_url = f"https://legalacts.ru{href}"
            await parse_document(session, full_url, existing_hashes)

    return True


async def main():
    existing_hashes = load_hashes(HASH_FILE)

    connector = TCPConnector(limit_per_host=1)
    async with ClientSession(headers=HEADERS, connector=connector) as session:
        page = 1
        while True:
            print(f"\n📄 Обработка страницы {page}")
            has_links = await parse_page(session, page, existing_hashes)
            if not has_links:
                print("🛑 Больше нет ссылок. Завершение.")
                break
            await asyncio.sleep(PAGE_DELAY)
            page += 1

    print("✅ Парсинг завершён.")


if __name__ == "__main__":
    asyncio.run(main())
